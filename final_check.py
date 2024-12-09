import json
import os
from docx import Document

enhanced_directory = 'enhanced_search_results/'
total_entries = 0
entries_with_downloads = 0
open_access_count = 0
open_access_downloaded = 0
entries_by_source = {}
downloaded_entries = []
non_downloaded_entries = []

for filename in os.listdir(enhanced_directory):
    if filename.endswith('.json'):
        filepath = os.path.join(enhanced_directory, filename)
        
        with open(filepath, 'r') as file:
            try:
                data = json.load(file)
                references = data.get('references', data) if isinstance(data, dict) else data
                if not isinstance(references, list):
                    print(f"Invalid format in: {filename}")
                    continue
                
                for entry in references:
                    total_entries += 1
                    is_open_access = entry.get("is_open_access", False)
                    if is_open_access:
                        open_access_count += 1
                    
                    if entry.get("downloaded_file"):
                        entries_with_downloads += 1
                        if is_open_access:
                            open_access_downloaded += 1
                        downloaded_entries.append(entry)
                        source = entry.get("download_source", "Unknown")
                        entries_by_source[source] = entries_by_source.get(source, 0) + 1
                    else:
                        non_downloaded_entries.append(entry)
                        
            except json.JSONDecodeError:
                print(f"Error decoding: {filename}")
                continue

download_percentage = (entries_with_downloads / total_entries * 100) if total_entries > 0 else 0
open_access_percentage = (open_access_count / total_entries * 100) if total_entries > 0 else 0
open_access_download_percentage = (open_access_downloaded / open_access_count * 100) if open_access_count > 0 else 0

doc = Document()
doc.add_heading('Download Analysis Report', level=1)

doc.add_paragraph(f"Total entries analyzed: {total_entries}")
doc.add_paragraph(f"Open access papers: {open_access_count} ({open_access_percentage:.2f}%)")
doc.add_paragraph(f"Open access papers downloaded: {open_access_downloaded} ({open_access_download_percentage:.2f}%)")
doc.add_paragraph(f"Total entries with successful downloads: {entries_with_downloads} ({download_percentage:.2f}%)")

doc.add_heading('Downloads by Source:', level=2)
for source, count in entries_by_source.items():
    percentage = (count / entries_with_downloads * 100) if entries_with_downloads > 0 else 0
    doc.add_paragraph(f"{source}: {count} ({percentage:.2f}%)")

doc.add_heading('Successfully Downloaded Papers:', level=2)
for entry in downloaded_entries:
    doc.add_paragraph(
        f"Title: {entry.get('title', 'Unknown Title')}\n"
        f"Authors: {', '.join(entry.get('authors', ['Unknown Author']))}\n"
        f"Year: {entry.get('year', 'Unknown Year')}\n"
        f"Source: {entry.get('download_source', 'Unknown Source')}\n"
        f"Open Access: {'Yes' if entry.get('is_open_access') else 'No'}\n"
    )

doc.add_heading('Papers Not Downloaded:', level=2)
for entry in non_downloaded_entries:
    doc.add_paragraph(
        f"Title: {entry.get('title', 'Unknown Title')}\n"
        f"Authors: {', '.join(entry.get('authors', ['Unknown Author']))}\n"
        f"Year: {entry.get('year', 'Unknown Year')}\n"
        f"Open Access: {'Yes' if entry.get('is_open_access') else 'No'}\n"
    )

print(f"\nDownload Analysis Summary:")
print(f"Total entries: {total_entries}")
print(f"Open access papers: {open_access_count} ({open_access_percentage:.2f}%)")
print(f"Open access papers downloaded: {open_access_downloaded} ({open_access_download_percentage:.2f}%)")
print(f"Successfully downloaded: {entries_with_downloads} ({download_percentage:.2f}%)")
print("\nDownloads by source:")
for source, count in entries_by_source.items():
    percentage = (count / entries_with_downloads * 100) if entries_with_downloads > 0 else 0
    print(f"{source}: {count} ({percentage:.2f}%)")

doc.save('Download_Analysis_Report.docx')